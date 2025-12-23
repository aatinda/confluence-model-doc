import os
import xml.etree.ElementTree as ET
import docx

from pprint import pprint

def get_namespaced_attribute(element, prefix_attr_name, ns_map):
    """
    Gets a namespaced attribute from an Element, expanding the namespace automatically.
    """
    parts = prefix_attr_name.split(':', 1)
    if len(parts) == 2:
        prefix, attr_name = parts
        if prefix in ns_map:
            expanded_attr_name = f"{{{ns_map[prefix]}}}{attr_name}"
            return element.attrib.get(expanded_attr_name)
    # If no prefix, or prefix not in map, try to get as a regular attribute
    return element.attrib.get(prefix_attr_name)

def return_property(properties, property_name):
    """
    Helper function to return a property value from our constructed page data.
    """

    value = None

    for prop in properties:
        if prop.get("name") == property_name:
            value = prop.get("value")
            break

    return value


def process_properties(uml_element, id_to_name_map=None):
    """
    Process all properties of an element.

    This will remove any with an "}" in the name which catches the xml:id and xml:type.
    """
    properties = []
    property = {}

    for prop_name, prop_value in uml_element.attrib.items():

        if '}' not in prop_name:

            if prop_name not in ['sType', 'nType', 'documentation']:
                # Resolve type to an actual data type name.
                if prop_name == 'type':
                    if id_to_name_map:
                        prop_value = id_to_name_map.get(prop_value, None)

                property = {
                    'name': prop_name,
                    'value': prop_value
                }

                properties.append(property)
    return properties


def generate_class_document(model_file, output_path):
    """
    Generate class documentation. One word doc for all classes.
    """
    tree = ET.parse(model_file)
    root = tree.getroot()

    ns = { 'xmi': 'http://schema.omg.org/spec/XMI/2.1', 'uml': 'http://schema.omg.org/spec/UML/2.0'}

    # Create word document for class documentation
    class_doc = docx.Document()

    # Find all Classes
    for packaged_element in root.findall('.//packagedElement[@xmi:type="uml:Class"]', ns):

        data = {}


        ###############################################################################
        ##  Find the element and extract properties.                                 ##
        ###############################################################################

        class_id = get_namespaced_attribute(packaged_element, "xmi:id", ns)
        xmi_type = get_namespaced_attribute(packaged_element, "xmi:type", ns)
        class_name = packaged_element.get('name')
        element = root.find(f'.//element[@xmi:idref="{class_id}"]', ns)

        class_desc = element.find('tags').find('tag[@name="definition"]').attrib.get('value')

        print(">> Processing class: ", class_name)

        ###############################################################################
        ##  Add main details of the class.                                           ##
        ###############################################################################

        data['details'] = {
            'id': class_id,
            'name': class_name,
            'type': xmi_type,
            'description': class_desc
        }

        ###############################################################################
        ##  Loop through the ownedAttributes.                                        ##
        ###############################################################################

        data['attributes'] = []
        attribute = {}

        for owned_attibute in packaged_element.findall('./ownedAttribute', ns):

            print("ownedAttribute Name: ", owned_attibute.get('name'))

            # Sparx Enterprise Architect puts associations as ownedAttribute of type uml:Property as well
            if not owned_attibute.get('association'):

                attr_id = get_namespaced_attribute(owned_attibute, 'xmi:id', ns)
                attr_name = owned_attibute.get('name')

                #######################################################################
                ##  Get the relevant attribute data.                                 ##
                #######################################################################

                attr_data = get_attribute_data(attr_id,
                                                    root,
                                                    ns)

                #######################################################################
                ##  Add attributes for class table.                                  ##
                #######################################################################

                attribute = {
                    'visibility': return_property(attr_data['properties'], 'scope'),
                    'name': attr_name,
                    'type': return_property(attr_data['properties'], 'type'),
                    'description': attr_data['details']['description'],
                    'optionality': [d for d in attr_data['properties'] if d.get("name") == "bounds"][0].get('value')
                }

                data['attributes'].append(attribute)

        ###############################################################################
        ##  Add heading, definition, and attribute table to the class documentation. ##
        ###############################################################################

        class_doc.add_heading(data['details']['name'])
        class_doc.add_paragraph(data['details']['description'])

        rows = [[attr["name"], attr["type"], attr["description"], attr["optionality"]] for attr in data['attributes']]
        class_table = class_doc.add_table(rows=1, cols=4, style="Table Grid")

        hdr_cells = class_table.rows[0].cells
        hdr_cells[0].text = "Attribute Name"
        hdr_cells[1].text = "Data Type"
        hdr_cells[2].text = "Definition"
        hdr_cells[3].text = "Optionality"

        for name, type, description, optionality in rows:
            row_cells = class_table.add_row().cells
            row_cells[0].text = str(name)
            row_cells[1].text = str(type)
            row_cells[2].text = str(description)
            row_cells[3].text = str(optionality)

    # Save Documentation
    os.makedirs(output_path, exist_ok=True)
    class_doc.save(os.path.join(output_path, 'classes.docx'))

def get_attribute_data(attr_id, root, ns):
    """
    Generate attribute documentation. One page per attribute.
    """

    print(">>>>>", attr_id)

    owned_attribute = root.find(f'.//ownedAttribute[@xmi:id="{attr_id}"]', ns)
    attribute = root.find(f'.//attribute[@xmi:idref="{attr_id}"]', ns)
    properties = attribute.find('properties')
    bounds = attribute.find('bounds')

    attr_data = {}

    attr_name = attribute.get('name')
    xmi_type = get_namespaced_attribute(owned_attribute, "xmi:type", ns)
    attr_desc = attribute.find('tags').find('tag[@name="definition"]').attrib.get('value')

    attr_data['details'] = {
        'id': attr_id,
        'name': attr_name,
        'type': xmi_type,
        'description': attr_desc
    }

    attr_data['properties'] = process_properties(properties)

    attr_data['properties'].append({ 'name': 'bounds', 'value': bounds.get('lower') + ".." + bounds.get('upper') })
    attr_data['properties'].append({ 'name': 'idref', 'value': attr_id })
    attr_data['properties'].append({ 'name': 'scope', 'value': attribute.get('scope') })

    pprint(attr_data)

    print("*****")

    return attr_data


def generate_enumeration_document(model_file, output_path):
    """
    Generate enumeration documentation. One page per enumeration.
    """
    tree = ET.parse(model_file)
    root = tree.getroot()

    ns = { 'xmi': 'http://schema.omg.org/spec/XMI/2.1', 'uml': 'http://schema.omg.org/spec/UML/2.0' }

    # Create Word document for enumeration documentation
    enum_doc = docx.Document()


    # Find all Enumerations
    for packaged_element in root.findall('.//packagedElement[@xmi:type="uml:Enumeration"]', ns):

        data = {}

        enum_id = get_namespaced_attribute(packaged_element, "xmi:id", ns)
        xmi_type = get_namespaced_attribute(packaged_element, "xmi:type", ns)
        enum_name = packaged_element.get('name')

        ###############################################################################
        ##  Find the element.                                                        ##
        ###############################################################################

        element = root.find(f'.//element[@xmi:idref="{enum_id}"]', ns)
        properties = element.find('properties')

        # enum_desc = properties.get('documentation')
        enum_desc = element.find('tags').find('tag[@name="definition"]').attrib.get('value')

        ###############################################################################
        ##  Add the main details of the enumeration.                                 ##
        ###############################################################################

        data['details'] = {
             'id': enum_id,
             'name': enum_name,
             'type': xmi_type,
             'description': enum_desc
        }

        ###############################################################################
        ##  Add the properties of the enumeration.                                   ##
        ###############################################################################

        data['properties'] = process_properties(properties)

        ###############################################################################
        ##  Loop through the ownedLiterals.                                          ##
        ###############################################################################

        data['literals'] = []
        literal = {}

        for owned_literal in packaged_element.findall('./ownedLiteral', ns):

            lit_id = get_namespaced_attribute(owned_literal, 'xmi:id', ns)
            lit_name = owned_literal.get('name')
            lit_type = owned_literal.get('type')

            ###########################################################################
            ##  Get the literal attributes.                                          ##
            ###########################################################################

            lit_data = get_literal_data(lit_id,
                                             root,
                                             ns,
                                             os.path.join(output_path, enum_name))

            ###########################################################################
            ##  End of get the literal attributes.                                   ##
            ###########################################################################

            literal = {
                'visibility': return_property(lit_data['properties'], 'scope'),
                'name': lit_name,
                'description': lit_data['details']['description']
            }

            data['literals'].append(literal)

        ###############################################################################
        ##  Add heading, definition, and literal table to the enum documentation.    ##
        ###############################################################################

        enum_doc.add_heading(data['details']['name'])
        enum_doc.add_paragraph(data['details']['description'])

        rows = [[lit["name"], lit["description"]] for lit in data['literals']]
        enum_table = enum_doc.add_table(rows=1, cols=2, style="Table Grid")

        hdr_cells = enum_table.rows[0].cells
        hdr_cells[0].text = "Value"
        hdr_cells[1].text = "Definition"

        for value, description in rows:
            row_cells = enum_table.add_row().cells
            row_cells[0].text = str(value)
            row_cells[1].text = str(description)

    os.makedirs(output_path, exist_ok=True)
    enum_doc.save(os.path.join(output_path, 'enumerations.docx'))



def get_literal_data(lit_id, root, ns, output_path):
    """
    Generate literal documentaiton. One page per literal.

    This only get the literal attributes currently. No page is generated.
    """

    owned_literal = root.find(f'.//ownedLiteral[@xmi:id="{lit_id}"]', ns)
    literal = root.find(f'.//attribute[@xmi:idref="{lit_id}"]', ns)
    documentation = literal.find('documentation')
    properties = literal.find('properties')
    bounds = literal.find('bounds')

    lit_data = {}

    lit_name = literal.get('name')
    xmi_type = get_namespaced_attribute(owned_literal, "xmi:type", ns)
    
    lit_desc = literal.find('tags').find('tag[@name="definition"]').attrib.get('value')

    lit_data['details'] = {
        'id': lit_id,
        'name': lit_name,
        'type': xmi_type,
        'description': lit_desc
    }

    lit_data['properties'] = process_properties(properties)

    lit_data['properties'].append({ 'name': 'bounds', 'value': bounds.get('lower') + ".." + bounds.get('upper') })
    lit_data['properties'].append({ 'name': 'idref', 'value': lit_id })
    lit_data['properties'].append({ 'name': 'scope', 'value': literal.get('scope') })

    return lit_data


def main():
    model_file = os.path.join("model", "TransportSafetyModel.xmi")
    output_dir = "output_word"

    generate_class_document(model_file, os.path.join(output_dir))

    generate_enumeration_document(model_file, os.path.join(output_dir))

if __name__ == "__main__":
    main()